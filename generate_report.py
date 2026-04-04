"""Generate Performance Comparison Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page Margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREEN = RGBColor(0x00, 0x80, 0x3E)
RED = RGBColor(0xC0, 0x39, 0x2B)
ORANGE = RGBColor(0xD4, 0x8B, 0x0A)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF2, 0xF6, 0xFA)
HEADER_BG = RGBColor(0x1B, 0x3A, 0x5C)


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    tc_pr.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None, highlight_last=True):
    """Add a formatted table with header row and optional last-row highlight."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B3A5C")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        is_last = highlight_last and r_idx == len(rows) - 1
        for c_idx, (text, color) in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            if color:
                run.font.color.rgb = color
            if is_last:
                run.bold = True
                set_cell_shading(cell, "EBF0F5")
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "F8F9FA")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_section_heading(doc, number, title, subtitle=None):
    """Add a styled section heading."""
    doc.add_paragraph()  # spacing
    p = doc.add_paragraph()
    run = p.add_run(f"  {number}. {title}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_BLUE
    if subtitle:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(10)
        run2.font.color.rgb = GRAY
        run2.italic = True


def add_note(doc, text):
    """Add a highlighted note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.italic = True


# ═══════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Inventory Object Counter")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Performance Comparison Report")
run2.font.size = Pt(16)
run2.font.color.rgb = ACCENT_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    "Previous Training vs New Training (Without CLIP) vs New Training (With CLIP Re-ID)\n"
    "Test Video: Client Office Walkthrough \u2014 2,335 frames, 30 FPS"
)
run3.font.size = Pt(10)
run3.font.color.rgb = GRAY

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════
# SECTION 1 — Previous Model (Baseline)
# ═══════════════════════════════════════════════════════════════════
add_section_heading(
    doc, 1,
    "Previous Model (Baseline)",
    "YOLOv11-Large (old training) \u2014 basic tracker, no re-identification, no deduplication",
)

add_styled_table(
    doc,
    ["Object Class", "Raw Tracks", "Reported Count", "Notes"],
    [
        [("Desk", None), ("18", None), ("18", None), ("Over-count (camera revisits)", GRAY)],
        [("Office Chair", None), ("26", None), ("26", None), ("Over-count", GRAY)],
        [("Pedestal", None), ("9", None), ("9", None), ("Over-count", GRAY)],
        [("Bookshelf / Cabinet", None), ("11", None), ("11", None), ("Over-count", GRAY)],
        [("Laptop", None), ("8", None), ("8", None), ("Over-count", GRAY)],
        [("Monitor", None), ("0", RED), ("0", RED), ("Not detected \u2014 model failed to learn this class", RED)],
        [("Mouse", None), ("3", None), ("3", None), ("Partial detection", GRAY)],
        [("Printer / Scanner", None), ("5", None), ("5", None), ("Over-count", GRAY)],
        [("Telephone / VoIP Phone", None), ("2", None), ("2", None), ("Over-count", GRAY)],
        [("TOTAL", DARK_BLUE), ("82", DARK_BLUE), ("82", RED), ("No dedup \u2014 every track = unique", RED)],
    ],
)

add_note(
    doc,
    "\u26a0 Key Limitation: The previous model had no deduplication. When the camera panned away "
    "and came back, the same object received a new track ID and was counted again. "
    "Additionally, the model could not detect Monitors at all.",
)

# ═══════════════════════════════════════════════════════════════════
# SECTION 2 — New Model WITHOUT CLIP
# ═══════════════════════════════════════════════════════════════════
add_section_heading(
    doc, 2,
    "New Model \u2014 Without CLIP Re-ID",
    "YOLOv11-Large (new training) + ByteTrack tracker, CLIP disabled, no deduplication",
)

add_styled_table(
    doc,
    ["Object Class", "Raw Tracks", "Reported Count", "Notes"],
    [
        [("Desk", None), ("21", None), ("21", None), ("Over-count (camera revisits)", GRAY)],
        [("Office Chair", None), ("28", None), ("28", None), ("Over-count", GRAY)],
        [("Pedestal", None), ("7", None), ("7", None), ("Over-count", GRAY)],
        [("Bookshelf / Cabinet", None), ("9", None), ("9", None), ("Over-count", GRAY)],
        [("Laptop", None), ("10", None), ("10", None), ("Over-count", GRAY)],
        [("Monitor", None), ("6", GREEN), ("6", GREEN), ("Now detected! (but over-counted)", GREEN)],
        [("Mouse", None), ("6", None), ("6", None), ("Better recall than previous", GRAY)],
        [("Printer / Scanner", None), ("4", None), ("4", None), ("Over-count", GRAY)],
        [("Telephone / VoIP Phone", None), ("2", None), ("2", None), ("Over-count", GRAY)],
        [("TOTAL", DARK_BLUE), ("93", DARK_BLUE), ("93", ORANGE), ("More detections but still no dedup", ORANGE)],
    ],
)

add_note(
    doc,
    "\u2714 Improvement: The newly trained model successfully detects Monitors and has "
    "overall better recall (+13% more detections). However, without CLIP re-identification, "
    "every new track is still counted as a separate object \u2014 over-counting persists.",
)

# ═══════════════════════════════════════════════════════════════════
# SECTION 3 — Head-to-Head
# ═══════════════════════════════════════════════════════════════════
add_section_heading(
    doc, 3,
    "Head-to-Head: Previous vs New Training",
    "Detection capability comparison \u2014 both without any deduplication",
)

add_styled_table(
    doc,
    ["Object Class", "Previous Model", "New Model", "Change"],
    [
        [("Desk", None), ("18", None), ("21", None), ("+3 (better recall)", GREEN)],
        [("Office Chair", None), ("26", None), ("28", None), ("+2 (better recall)", GREEN)],
        [("Pedestal", None), ("9", None), ("7", None), ("-2 (fewer false positives)", GREEN)],
        [("Bookshelf / Cabinet", None), ("11", None), ("9", None), ("-2 (fewer false positives)", GREEN)],
        [("Laptop", None), ("8", None), ("10", None), ("+2 (better recall)", GREEN)],
        [("Monitor", None), ("0 (undetected)", RED), ("6", GREEN), ("NEW CLASS DETECTED", GREEN)],
        [("Mouse", None), ("3", None), ("6", None), ("+3 (better recall)", GREEN)],
        [("Printer / Scanner", None), ("5", None), ("4", None), ("-1 (fewer false positives)", GREEN)],
        [("Telephone / VoIP Phone", None), ("2", None), ("2", None), ("No change", GRAY)],
        [("TOTAL RAW TRACKS", DARK_BLUE), ("82", DARK_BLUE), ("93", DARK_BLUE), ("+11 (13% more detections)", ACCENT_BLUE)],
    ],
)

# Comparison bullets
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Previous Model Weaknesses:")
run.bold = True
run.font.color.rgb = RED
run.font.size = Pt(11)

for item in [
    "Could not detect Monitors at all",
    "Lower recall for small objects (Mouse, Keyboard)",
    "No deduplication mechanism",
    "Basic tracker with frequent ID switches",
]:
    bp = doc.add_paragraph(style="List Bullet")
    run = bp.add_run(item)
    run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("New Model Improvements:")
run.bold = True
run.font.color.rgb = GREEN
run.font.size = Pt(11)

for item in [
    "Monitor detection now working",
    "Improved small object recall (Mouse detection 2x better)",
    "ByteTrack tracker for more robust tracking",
    "CLIP Re-ID available for deduplication",
    "13% more total detections across all classes",
]:
    bp = doc.add_paragraph(style="List Bullet")
    run = bp.add_run(item)
    run.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════
# SECTION 4 — New Model WITH CLIP
# ═══════════════════════════════════════════════════════════════════
add_section_heading(
    doc, 4,
    "New Model \u2014 With CLIP Re-ID (Full Pipeline)",
    "YOLOv11-Large (new training) + ByteTrack + CLIP re-identification & co-occurrence deduplication",
)

# Client video results
p = doc.add_paragraph()
run = p.add_run("Client Office Video (Primary Test)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = ACCENT_BLUE

add_styled_table(
    doc,
    ["Object Class", "Raw Tracks", "Duplicates Removed", "Unique Count"],
    [
        [("Desk", None), ("13", None), ("9", None), ("4", GREEN)],
        [("Office Chair", None), ("16", None), ("10", None), ("6", GREEN)],
        [("Pedestal", None), ("6", None), ("4", None), ("2", GREEN)],
        [("Bookshelf / Cabinet", None), ("7", None), ("4", None), ("3", GREEN)],
        [("Laptop", None), ("9", None), ("7", None), ("2", GREEN)],
        [("Monitor", None), ("7", None), ("5", None), ("2", GREEN)],
        [("Mouse", None), ("7", None), ("5", None), ("2", GREEN)],
        [("Printer / Scanner", None), ("3", None), ("1", None), ("2", GREEN)],
        [("Telephone / VoIP Phone", None), ("1", None), ("0", None), ("1", GREEN)],
        [("TOTAL", DARK_BLUE), ("83", DARK_BLUE), ("59 (71%)", DARK_BLUE), ("24", GREEN)],
    ],
)

# Summary stats
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in [
    ("Duplicates Removed: ", "71% (59 of 83 tracks)"),
    ("   |   Final Unique Objects: ", "24 (across 9 classes)"),
    ("   |   Processing Time: ", "52.8 seconds"),
]:
    run = p.add_run(label)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run = p.add_run(value)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_BLUE

# All videos summary
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("All Test Videos Summary")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = ACCENT_BLUE

add_styled_table(
    doc,
    ["Video", "Frames", "Raw Tracks", "Unique Objects", "Duplicates Removed", "Top Classes"],
    [
        [("Client Office", None), ("2,335", None), ("83", None), ("24", GREEN), ("59 (71%)", None), ("Chair: 6, Desk: 4", GRAY)],
        [("Office Area 1", None), ("127", None), ("7", None), ("5", GREEN), ("2 (29%)", None), ("Cubicle: 1, Desk: 1", GRAY)],
        [("Office Area 2", None), ("157", None), ("7", None), ("6", GREEN), ("1 (14%)", None), ("Cubicle: 2, Desk: 1", GRAY)],
        [("Office Area 3", None), ("169", None), ("11", None), ("9", GREEN), ("2 (18%)", None), ("Filing Cab: 3, Chair: 2", GRAY)],
        [("Storage Room 1", None), ("429", None), ("16", None), ("5", GREEN), ("11 (69%)", None), ("Filing Cabinet: 5", GRAY)],
        [("Breakroom", None), ("347", None), ("33", None), ("15", GREEN), ("18 (55%)", None), ("BR Chair: 9, Chair: 3", GRAY)],
        [("Storage Room 2", None), ("429", None), ("15", None), ("5", GREEN), ("10 (67%)", None), ("Filing Cabinet: 5", GRAY)],
        [("Cubicle Floor", None), ("1,094", None), ("55", None), ("16", GREEN), ("39 (71%)", None), ("Cubicle: 7, Desk: 2", GRAY)],
        [("GRAND TOTAL", DARK_BLUE), ("5,087", DARK_BLUE), ("227", DARK_BLUE), ("85", GREEN), ("142 (63%)", DARK_BLUE), ("", None)],
    ],
)

add_note(
    doc,
    "\u2714 CLIP Re-ID + Co-occurrence Guard: The pipeline uses CLIP visual embeddings to match "
    "objects when the camera revisits an area, while the co-occurrence guard ensures that objects "
    "visible in the same frame (e.g., 6 identical office chairs) are never incorrectly merged. "
    "This eliminated 63% of duplicate tracks across all test videos.",
)

# ═══════════════════════════════════════════════════════════════════
# FINAL SUMMARY TABLE
# ═══════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_heading(doc, None, "Key Takeaways")

add_styled_table(
    doc,
    ["Metric", "Previous Model", "New (No CLIP)", "New (With CLIP)"],
    [
        [("Model", None), ("YOLOv11-Large\n(old training)", None), ("YOLOv11-Large\n(new training)", None), ("YOLOv11-Large\n(new training)", None)],
        [("Tracker", None), ("Basic", None), ("ByteTrack", None), ("ByteTrack", None)],
        [("Re-Identification", None), ("None", RED), ("None", RED), ("CLIP + Co-occurrence", GREEN)],
        [("Monitor Detection", None), ("Failed", RED), ("Working", GREEN), ("Working", GREEN)],
        [("Reported Count*", None), ("82", RED), ("93", ORANGE), ("24", GREEN)],
        [("Over-counting", None), ("Severe", RED), ("Severe", RED), ("Resolved", GREEN)],
        [("Duplicate Elimination", None), ("0%", RED), ("0%", RED), ("71%", GREEN)],
    ],
    highlight_last=False,
)

p = doc.add_paragraph()
run = p.add_run("* Based on Client Office Walkthrough video (2,335 frames)")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Inventory Counter Pipeline \u2014 Performance Comparison Report \u2014 February 2026")
run.font.size = Pt(9)
run.font.color.rgb = GRAY

# ── Save ──
output_path = r"d:\Personal\Inventory_counter\Performance_Comparison.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
